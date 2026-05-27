"""Append simple Chapter 4 and Chapter 5 content to the original .docx and save a new file.

Usage: run from workspace root: `python -m scripts.add_chapters` or `python scripts/add_chapters.py`
"""
from docx import Document
from pathlib import Path


SRC = Path.home() / "Downloads" / "Privacy Sentinel FINAL v3.docx"
OUT = Path("docs") / "Privacy_Sentinel_FINAL_v3_with_ch4_5.docx"


CH4 = (
    "Chapter 4: Implementation (Added)\n\n"
    "This implementation provides a minimal, reproducible Python implementation of the core Privacy Sentinel components:\n"
    "- SQLite database initialisation (data/privacy_sentinel.db).\n"
    "- A comparator for simple set-based comparisons of declared data categories and data types.\n"
    "- A lightweight scoring engine producing completeness, specificity and consistency sub-scores.\n"
    "- A minimal Streamlit UI for quick, human-driven comparisons and visualisation.\n"
)

CH5 = (
    "Chapter 5: Evaluation and Conclusion (Added)\n\n"
    "This simplified implementation is intended as a demonstrator rather than a production tool.\n"
    "Evaluation steps include: running the built-in tests with pytest, interacting with the Streamlit UI,\n"
    "and running comparisons on a small sample of apps. The conclusion highlights limitations (no APK analysis,\n"
    "limited automated verification) and suggests future work: runtime instrumentation, larger datasets,\n"
    "and improved scoring logic.\n"
)


def main():
    if not SRC.exists():
        print(f"Source file not found: {SRC}")
        return
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(SRC))

    # Append chapter 4
    doc.add_page_break()
    doc.add_heading("Chapter 4: Implementation", level=1)
    for line in CH4.split("\n"):
        doc.add_paragraph(line)

    # Append chapter 5
    doc.add_page_break()
    doc.add_heading("Chapter 5: Evaluation and Conclusion", level=1)
    for line in CH5.split("\n"):
        doc.add_paragraph(line)

    doc.save(str(OUT))
    print("Saved updated document:", OUT)


if __name__ == "__main__":
    main()
